import json
import random
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE_DOC = Path(r"G:\秒懂词更新\题库\初一\5.24初一更新题库.docx")
OUTPUT_DOCX = Path(r"G:\秒懂词更新\题库\初一\5.24初一更新题库_全量句子选择题.docx")
OUTPUT_JSON = Path(r"G:\秒懂词更新\题库\初一\5.24初一更新题库_全量句子选择题.json")
PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_JS = PROJECT_ROOT / "word-data" / "quiz-grade7-sentences.js"
REPORT_JSON = Path(r"G:\秒懂词更新\题库\初一\5.24初一更新题库_全量句子选择题_解析报告.json")
EXTERNAL_VOCAB = Path(
    r"D:\Users\Administrator\Documents\xwechat_files\wxid_grs0vsuirik322_f31f\temp\RWTemp\2026-05\6cb53c9d512118515fb069877ce75d68\初一词库(2).txt"
)
PROJECT_VOCAB = PROJECT_ROOT / "word-data" / "初一词库.txt"


SECTION_HEADING_RE = re.compile(
    r"^(?:[一二三四五六七八九十]+[.．、]|（[ABC]）|\([ABC]\)|用括号内所给单词|（C）|\(C\))"
)
NUMBERED_PREFIX_RE = re.compile(r"^\s*(\d+)[.．]\s*")
PAREN_PROMPT_RE = re.compile(r"(?<!\d)(\d+)\s*[（(][^()（）]+[）)]")
STANDALONE_NUM_RE = re.compile(r"(?<![\d:,])\b([1-5])\b(?![\d:,])")
ANSWER_MARKER_RE = re.compile(r"(\d+)[.．]\s*")
SENTENCE_SPLIT_RE = re.compile(r"(?<=[.!?。！？])\s+")


@dataclass
class Block:
    kind: str
    text: str


@dataclass
class QuestionGroup:
    label: str
    source_type: str
    prompt_blocks: list[str] = field(default_factory=list)
    word_bank: list[str] = field(default_factory=list)
    answer_tokens: list[str] = field(default_factory=list)
    start_index: int = 0

    def target_numbers(self) -> list[int]:
        nums: list[int] = []
        for text in self.prompt_blocks:
            for n in find_target_numbers(text, self.source_type):
                nums.append(n)
        return nums


def normalize_text(text: str) -> str:
    return " ".join(text.replace("\u3000", " ").split())


def load_reference_vocab(paths: list[Path] | None = None) -> list[str]:
    paths = paths or [EXTERNAL_VOCAB, PROJECT_VOCAB]
    vocab: list[str] = []
    seen: set[str] = set()
    for path in paths:
        if not path.exists():
            continue
        for raw in path.read_text(encoding="utf-8").splitlines():
            line = normalize_text(raw)
            if not line or line.lower().startswith("unit "):
                continue
            english = re.split(r"(?=[\u4e00-\u9fff（(])", line, maxsplit=1)[0].strip(" ,，;；")
            if not english or not re.search(r"[A-Za-z]", english):
                continue
            key = english.lower()
            if key not in seen:
                seen.add(key)
                vocab.append(english)
    return vocab


def iter_blocks(doc: Document) -> Iterable[Block]:
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            texts = [node.text for node in child.iter(qn("w:t")) if node.text]
            text = normalize_text("".join(texts))
            if text:
                yield Block("paragraph", text)
        elif child.tag == qn("w:tbl"):
            rows = []
            for row in child.iter(qn("w:tr")):
                cells = []
                for cell in row.iter(qn("w:tc")):
                    texts = [node.text for node in cell.iter(qn("w:t")) if node.text]
                    cell_text = normalize_text("".join(texts))
                    if cell_text:
                        cells.append(cell_text)
                if cells:
                    rows.append(" ".join(cells))
            text = normalize_text(" ".join(rows))
            if text:
                yield Block("table", text)


def infer_source_type(text: str) -> str:
    if "汉语提示" in text:
        return "汉语提示"
    if "括号" in text or "适当形式" in text:
        return "适当形式"
    if "方框" in text:
        return "方框选词"
    if "首字母" in text:
        return "首字母"
    return "句子填空"


def is_question_heading(text: str) -> bool:
    return bool(SECTION_HEADING_RE.search(text))


def split_word_bank(text: str) -> list[str]:
    text = text.replace("succeedbuild", "succeed build")
    parts = re.split(r"[,，;；\s]+", text)
    words: list[str] = []
    phrase = []
    # Tables are one-line word banks; keep common multi-word phrases by splitting on
    # larger gaps is not available after DOCX extraction, so preserve known phrase spans.
    known_phrases = [
        "little by little",
        "look after",
        "after all",
        "keep fit",
        "go with",
        "jump rope",
        "be late for",
        "put on weight",
        "focus on",
        "good eating habits",
        "full of energy",
        "a few hours",
        "set off",
    ]
    lower = f" {text.lower()} "
    for phrase_text in known_phrases:
        if f" {phrase_text} " in lower and phrase_text not in words:
            words.append(phrase_text)
    for part in parts:
        cleaned = clean_answer(part)
        if cleaned and cleaned.lower() not in {w.lower() for w in words}:
            words.append(cleaned)
    return words


def looks_like_answer_line(text: str) -> bool:
    if "【答案】" in text:
        return True
    tokens = extract_answer_tokens(text)
    if len(tokens) < 3:
        return False
    if any(ch in text for ch in "()（）"):
        return False
    marker_count = len(ANSWER_MARKER_RE.findall(text))
    return marker_count >= 3


def clean_answer(answer: str) -> str:
    answer = normalize_text(answer)
    answer = re.sub(r"^【答案】", "", answer)
    answer = re.sub(r"^答案[:：]?", "", answer)
    answer = re.sub(r"^故填", "", answer)
    answer = answer.strip(" 。；;，,、")
    return answer


def extract_answer_tokens(text: str) -> list[str]:
    text = normalize_text(text)
    text = text.replace("【答案】", " ")
    text = re.sub(r"(?<=[A-Za-z/)）])(?=\d+[.．])", " ", text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])(?=\d+[.．])", " ", text)
    if not text:
        return []

    leading = ""
    first_marker = ANSWER_MARKER_RE.search(text)
    if first_marker and first_marker.group(1) != "1":
        leading = text[: first_marker.start()].strip()
        text = "1." + leading + " " + text[first_marker.start() :]

    matches = list(ANSWER_MARKER_RE.finditer(text))
    if not matches:
        return [clean_answer(text)] if clean_answer(text) else []

    tokens: list[str] = []
    for idx, match in enumerate(matches):
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        token = clean_answer(text[start:end])
        if token:
            tokens.append(token)
    return tokens


def find_target_numbers(text: str, source_type: str = "") -> list[int]:
    prefix = NUMBERED_PREFIX_RE.search(text)
    if prefix and len(PAREN_PROMPT_RE.findall(text)) <= 1 and source_type not in {"方框选词"}:
        return [int(prefix.group(1))]

    nums = [int(n) for n in PAREN_PROMPT_RE.findall(text)]
    if source_type == "方框选词" or not nums:
        for n in STANDALONE_NUM_RE.findall(text):
            value = int(n)
            if value not in nums:
                nums.append(value)
    return nums


def is_prompt_text(text: str, source_type: str) -> bool:
    if "【答案】" in text or "答案与解析" in text or "句意：" in text:
        return False
    return bool(find_target_numbers(text, source_type))


def collect_question_groups(blocks: list[Block]) -> tuple[list[QuestionGroup], list[list[str]], dict]:
    groups: list[QuestionGroup] = []
    answer_blocks: list[list[str]] = []
    current: QuestionGroup | None = None
    in_answer_section = False
    stats = {
        "nonempty_blocks": len(blocks),
        "tables": sum(1 for b in blocks if b.kind == "table"),
        "short_text_sections": 0,
        "answer_lines": 0,
    }

    def flush_current():
        nonlocal current
        if current and current.prompt_blocks:
            groups.append(current)
        current = None

    pending_answer: list[str] = []
    collecting_answer = False
    for idx, block in enumerate(blocks):
        text = block.text
        if "答案与解析" in text:
            flush_current()
            in_answer_section = True
            continue

        if in_answer_section:
            if "【答案】" in text:
                if pending_answer:
                    answer_blocks.append(pending_answer)
                    pending_answer = []
                collecting_answer = True
                tokens = extract_answer_tokens(text)
                if tokens:
                    pending_answer.extend(tokens)
                stats["answer_lines"] += 1
                continue
            if collecting_answer:
                if "【导语】" in text or "句意：" in text or is_question_heading(text):
                    if pending_answer:
                        answer_blocks.append(pending_answer)
                    pending_answer = []
                    collecting_answer = False
                    continue
                if re.fullmatch(r"\d{2}", text) or text in {"基础篇", "提升篇", "综合篇"}:
                    answer_blocks.append(pending_answer)
                    pending_answer = []
                    collecting_answer = False
                    continue
                tokens = extract_answer_tokens(text)
                if tokens:
                    pending_answer.extend(tokens)
                    stats["answer_lines"] += 1
                else:
                    if pending_answer:
                        answer_blocks.append(pending_answer)
                    pending_answer = []
                    collecting_answer = False
                continue
            continue

        if block.kind == "table":
            if current:
                current.word_bank.extend(split_word_bank(text))
            continue

        if is_question_heading(text):
            flush_current()
            source_type = infer_source_type(text)
            current = QuestionGroup(label=text, source_type=source_type, start_index=idx)
            if "短文" in text:
                stats["short_text_sections"] += 1
            continue

        if looks_like_answer_line(text) and current and current.prompt_blocks:
            current.answer_tokens = extract_answer_tokens(text)
            continue

        if current is None:
            current = QuestionGroup(label="未命名题组", source_type="句子填空", start_index=idx)

        if is_prompt_text(text, current.source_type):
            current.prompt_blocks.append(text)

    if pending_answer:
        answer_blocks.append(pending_answer)
    flush_current()
    return groups, answer_blocks, stats


def map_group_answers(groups: list[QuestionGroup], answer_blocks: list[list[str]]) -> tuple[dict[int, dict[int, str]], list[str]]:
    answer_maps: dict[int, dict[int, str]] = {}
    unresolved: list[str] = []
    answer_index = 0
    for group_index, group in enumerate(groups):
        targets = group.target_numbers()
        tokens = group.answer_tokens
        if not tokens and answer_index < len(answer_blocks):
            tokens = answer_blocks[answer_index]
            answer_index += 1
        if len(tokens) < len(targets):
            unresolved.append(
                f"{group.label}: 目标空 {len(targets)} 个，但答案 {len(tokens)} 个"
            )
            continue
        answer_maps[group_index] = {
            target: clean_answer(tokens[idx]) for idx, target in enumerate(targets)
        }
    return answer_maps, unresolved


def replace_numbered_prompts(text: str, target_number: int, answers: dict[int, str]) -> str:
    def paren_repl(match):
        number = int(match.group(1))
        return "________" if number == target_number else answers.get(number, "________")

    result = PAREN_PROMPT_RE.sub(paren_repl, text)

    def standalone_repl(match):
        number = int(match.group(1))
        if number not in answers and number != target_number:
            return match.group(0)
        return "________" if number == target_number else answers.get(number, match.group(0))

    result = STANDALONE_NUM_RE.sub(standalone_repl, result)
    return result


def replace_first_letter_prompt(text: str, answer: str) -> str:
    first = re.escape(answer[0])
    patterns = [
        rf"\b{first}\s+([,.;!?])",
        rf"\b{first}\s+(?=[a-zA-Z])",
        rf"\b{first}_+\b",
    ]
    for pattern in patterns:
        new_text, count = re.subn(pattern, lambda m: "________" + (m.group(1) if m.groups() else " "), text, count=1, flags=re.I)
        if count:
            return normalize_text(new_text)
    return text


def replace_first_parenthetical_prompt(text: str) -> str:
    return re.sub(r"\s*[（(][^()（）]+[）)]", " ________", text, count=1)


def select_sentence_with_blank(text: str) -> str:
    parts = SENTENCE_SPLIT_RE.split(text)
    for part in parts:
        if "________" in part:
            return normalize_text(part)
    return normalize_text(text)


def clean_sentence_for_single_blank(text: str, target_number: int, answers: dict[int, str]) -> str:
    text = NUMBERED_PREFIX_RE.sub("", text)
    result = replace_numbered_prompts(text, target_number, answers)
    if "________" not in result:
        result = replace_first_parenthetical_prompt(result)
    if "________" not in result and target_number in answers:
        result = replace_first_letter_prompt(result, answers[target_number])
    sentence = select_sentence_with_blank(result)
    sentence = re.sub(r"\s+([,.;!?])", r"\1", sentence)
    sentence = re.sub(r"\s{2,}", " ", sentence)
    return sentence.strip()


def answer_family(answer: str) -> list[str]:
    if "/" in answer:
        return []
    lower = answer.lower()
    variants = set()
    if " " in answer:
        parts = answer.split()
        if parts[0].lower() in {"am", "is", "are", "was", "were", "be", "being", "been"} and len(parts) == 2:
            verb = parts[1]
            variants.update([verb, verb.rstrip("s"), verb + "s", "is " + verb, "are " + verb, "was " + verb])
        elif parts[0].lower() == "to" and len(parts) == 2:
            verb = parts[1]
            variants.update([verb, verb + "s", verb + "ing", verb + "ed"])
        return [v for v in variants if v.lower() != lower and len(v) > 1]
    if lower.endswith("ies"):
        base = answer[:-3] + "y"
        variants.update([base, base + "s"])
    elif lower.endswith("ves"):
        base = answer[:-3] + "f"
        variants.update([base, answer[:-3] + "fe"])
    elif lower.endswith("ing"):
        base = answer[:-3]
        variants.update([base, base + "s", base + "ed"])
        if len(base) > 1 and base[-1] == base[-2]:
            variants.add(base[:-1])
    elif lower.endswith("ed"):
        base = answer[:-2]
        variants.update([base, base + "s", base + "ing"])
    elif lower.endswith("es"):
        base = answer[:-2]
        variants.update([base, base + "s", base + "ing"])
    elif lower.endswith("s") and len(answer) > 3:
        base = answer[:-1]
        variants.update([base, base + "ing", base + "ed"])
    else:
        if lower.endswith("y") and len(answer) > 2:
            variants.update([answer[:-1] + "ier", answer[:-1] + "iest", answer[:-1] + "ily"])
            return [v for v in variants if v.lower() != lower and len(v) > 1]
        variants.add(answer + "s")
    return [v for v in variants if v.lower() != lower and len(v) > 1]


def build_options(answer: str, answer_pool: list[str], word_bank: list[str], seed: str) -> list[str]:
    answer = clean_answer(answer)
    used = {answer.lower()}
    candidates: list[str] = []

    def add(candidate: str):
        candidate = clean_answer(candidate)
        if not candidate:
            return
        key = candidate.lower()
        if key not in used:
            used.add(key)
            candidates.append(candidate)

    first = answer[0].lower()
    combined_pool = word_bank + answer_pool
    if " " in answer:
        first_word = answer.split()[0].lower()
        for item in combined_pool:
            if item.lower().split()[0] == first_word:
                add(item)

    for item in combined_pool:
        if item and item[0].lower() == first and (" " in item) == (" " in answer):
            add(item)

    for variant in answer_family(answer):
        add(variant)

    for item in combined_pool:
        if item and item[0].lower() == first:
            add(item)

    for item in combined_pool:
        add(item)

    while len(candidates) < 4:
        add(f"{answer}{len(candidates) + 1}")

    options = [answer] + candidates[:4]
    rng = random.Random(seed)
    rng.shuffle(options)
    return options


def build_questions(
    groups: list[QuestionGroup],
    answer_maps: dict[int, dict[int, str]],
    reference_vocab: list[str] | None = None,
) -> tuple[list[dict], list[dict]]:
    answer_pool = list(reference_vocab or [])
    for answers in answer_maps.values():
        answer_pool.extend(answers.values())

    questions: list[dict] = []
    review: list[dict] = []
    counter = 1
    for group_index, group in enumerate(groups):
        answers = answer_maps.get(group_index)
        if not answers:
            review.append({"source": group.label, "reason": "缺少答案映射"})
            continue
        for block in group.prompt_blocks:
            for target in find_target_numbers(block, group.source_type):
                answer = answers.get(target)
                if not answer:
                    review.append({"source": group.label, "text": block, "target": target, "reason": "缺少目标答案"})
                    continue
                sentence = clean_sentence_for_single_blank(block, target, answers)
                if sentence.count("________") != 1:
                    review.append({"source": group.label, "text": block, "target": target, "reason": "无法拆成单空句"})
                    continue
                qid = f"g7-20260524-{counter:04d}"
                options = build_options(answer, answer_pool, group.word_bank, qid)
                questions.append(
                    {
                        "id": qid,
                        "grade": "初一",
                        "sentence": sentence,
                        "answer": answer,
                        "options": options,
                        "sourceType": group.source_type,
                        "sourceLabel": group.label,
                    }
                )
                counter += 1
    return questions, review


def validate_questions(questions: list[dict]) -> list[str]:
    errors = []
    for q in questions:
        if q["sentence"].count("________") != 1:
            errors.append(f"{q['id']} blank count invalid")
        if not q["answer"]:
            errors.append(f"{q['id']} missing answer")
        if len(q["options"]) != 5:
            errors.append(f"{q['id']} option count invalid")
        if q["answer"] not in q["options"]:
            errors.append(f"{q['id']} answer not in options")
        if len({o.lower() for o in q["options"]}) != 5:
            errors.append(f"{q['id']} duplicate options")
    return errors


def write_json(questions: list[dict]):
    OUTPUT_JSON.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_JSON.write_text(json.dumps(questions, ensure_ascii=False, indent=2), encoding="utf-8")


def write_js(questions: list[dict]):
    data = json.dumps(questions, ensure_ascii=False, indent=2)
    OUTPUT_JS.write_text(f"window.WORD_SNAP_GRADE7_QUIZ_SENTENCES = {data};\n", encoding="utf-8")


def set_run_font(run, size=9, bold=False):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold


def write_review_docx(questions: list[dict], review: list[dict], stats: dict):
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(9)

    title = doc.add_paragraph()
    run = title.add_run("5.24初一更新题库 - 网页导入句子选择题")
    set_run_font(run, size=16, bold=True)
    doc.add_paragraph(f"共 {len(questions)} 题。每题为单句填空，含 1 个正确答案和 4 个干扰选项。")

    headers = ["题号", "句子题目", "正确答案", "A", "B", "C", "D", "E"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for idx, q in enumerate(questions, 1):
        row = table.add_row().cells
        values = [str(idx), q["sentence"], q["answer"], *q["options"]]
        for cell, value in zip(row, values):
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    if review:
        doc.add_page_break()
        doc.add_heading("未进入网页题库的待确认题", level=1)
        review_table = doc.add_table(rows=1, cols=4)
        review_table.style = "Table Grid"
        for cell, header in zip(review_table.rows[0].cells, ["来源", "原文", "目标", "原因"]):
            cell.text = header
        for item in review:
            row = review_table.add_row().cells
            row[0].text = item.get("source", "")
            row[1].text = item.get("text", "")
            row[2].text = str(item.get("target", ""))
            row[3].text = item.get("reason", "")

    doc.save(OUTPUT_DOCX)


def build_all(source_doc: Path = SOURCE_DOC) -> tuple[list[dict], list[dict], dict]:
    doc = Document(source_doc)
    blocks = list(iter_blocks(doc))
    groups, answer_blocks, stats = collect_question_groups(blocks)
    answer_maps, unresolved = map_group_answers(groups, answer_blocks)
    reference_vocab = load_reference_vocab()
    questions, review = build_questions(groups, answer_maps, reference_vocab)
    review.extend({"source": item, "reason": "答案数量不匹配"} for item in unresolved)
    errors = validate_questions(questions)
    if errors:
        raise ValueError("\n".join(errors[:20]))
    stats.update(
        {
            "groups": len(groups),
            "answer_blocks": len(answer_blocks),
            "questions": len(questions),
            "review": len(review),
            "raw_target_blanks": sum(len(group.target_numbers()) for group in groups),
            "reference_vocab": len(reference_vocab),
        }
    )
    return questions, review, stats


def main():
    questions, review, stats = build_all()
    write_json(questions)
    write_js(questions)
    write_review_docx(questions, review, stats)
    REPORT_JSON.write_text(
        json.dumps({"stats": stats, "review": review[:200]}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(json.dumps(stats, ensure_ascii=False, indent=2))
    print(OUTPUT_DOCX)
    print(OUTPUT_JSON)
    print(OUTPUT_JS)


if __name__ == "__main__":
    main()
